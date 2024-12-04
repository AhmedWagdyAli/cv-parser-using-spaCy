import sys
from cv_handler import CVHandler
from cv_parser import CVParser
from docx import Document
import os  # Add this line

class CVProcessor:
    """Central dispatcher for processing CVs."""
    def __init__(self):
        self.parser = CVParser()

    def process(self, file_path):
        try:
            cv_text = CVHandler.extract_text(file_path)
            if not cv_text.strip():
                print("No text extracted from the file.")
                return None
            return self.parser.parse(cv_text)
        except Exception as e:
            print(f"Error processing CV: {e}")
            return None
   
    @staticmethod
    def fill_template(parsed_data, template_path, output_path):
        """Fill a Word template with parsed CV data."""
        try:
            # Ensure the template exists
            if not os.path.isfile(template_path):
                raise FileNotFoundError(f"Template not found at: {template_path}")

            doc = Document(template_path)

            # Helper function to flatten experience data
            def format_experience(experience_list):
                if not isinstance(experience_list, list):
                    return "N/A"
                formatted_experience = []
                for item in experience_list:
                    details = []
                    for key, value in item.items():
                        if value:  # Include only non-empty values
                            details.append(f"{key.capitalize()}: {value}")
                    formatted_experience.append("\n".join(details))
                return "\n\n".join(formatted_experience)

            # Helper function to format skills as bullet points
            def format_skills(skills_text):
                if not skills_text or not isinstance(skills_text, str):
                    return "N/A"
                skills_list = [skill.strip() for skill in skills_text.split(",")]
                return "\n".join(f"• {skill}" for skill in skills_list if skill)

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in parsed_data.items():
                    placeholder = f"{{{{{key}}}}}"  # e.g., {{name}}

                    if key == "experience" and isinstance(value, list):
                        value = format_experience(value)
                    elif key == "skills" and isinstance(value, str):
                        value = format_skills(value)
                    elif isinstance(value, list):
                        value = "\n".join(value)
                    elif isinstance(value, dict):
                        value = "\n".join(f"{k}: {v}" for k, v in value.items() if v)
                    else:
                        value = str(value) if value else "N/A"

                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in parsed_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if key == "experience" and isinstance(value, list):
                                value = format_experience(value)
                            elif key == "skills" and isinstance(value, str):
                                value = format_skills(value)
                            elif isinstance(value, list):
                                value = "\n".join(value)
                            elif isinstance(value, dict):
                                value = "\n".join(f"{k}: {v}" for k, v in value.items() if v)
                            else:
                                value = str(value) if value else "N/A"

                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)

            # Save the filled template
            doc.save(output_path)
            print(f"Filled CV saved to: {output_path}")
        except Exception as e:
            print(f"Error filling the template: {e}")

if __name__ == "__main__":
    import sys

    # Ensure correct usage
    if len(sys.argv) < 3:
        print("Usage: python cv_processor.py <path_to_cv_file> <path_to_template>")
        sys.exit(1)

    # Extract command-line arguments
    file_path = sys.argv[1]  # Path to the CV file
    template_path = sys.argv[2]  # Path to the Word template file

    # Process the CV
    processor = CVProcessor()
    parsed_data = processor.process(file_path)

    if parsed_data:
        print("Parsed Data:")
        print(parsed_data)

        # Fill the template with parsed data
        output_file = "filled_cv.docx"  # Define the output file name
        processor.fill_template(parsed_data, template_path=template_path, output_path=output_file)
    else:
        print("Failed to process the CV.")
